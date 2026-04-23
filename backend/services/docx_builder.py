"""Render canonical CV JSON to a DOCX, in three template styles.

Rather than depending on the fragile {{placeholder}} system from the old app
(which required hand-edited Word templates with invisible formatting tags),
we build a clean DOCX from scratch. This gives pixel-predictable output and
makes adding a template = adding a small style config.
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm


# --- Style presets -------------------------------------------------------------
STYLES = {
    "minimal": {
        "font": "Calibri",
        "name_size": 26,
        "accent": RGBColor(0x1F, 0x2D, 0x3D),
        "rule": True,
    },
    "classic": {
        "font": "Garamond",
        "name_size": 24,
        "accent": RGBColor(0x12, 0x2E, 0x5B),
        "rule": True,
    },
    "modern": {
        "font": "Inter",
        "name_size": 28,
        "accent": RGBColor(0x25, 0x63, 0xEB),
        "rule": False,
    },
}


def _heading(doc: Document, text: str, style: dict):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = style["font"]
    run.font.color.rgb = style["accent"]
    if style["rule"]:
        # Use an em-dash rule under section headings
        rule = doc.add_paragraph()
        rule.paragraph_format.space_before = Pt(0)
        rule.paragraph_format.space_after = Pt(4)
        r = rule.add_run("─" * 60)
        r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        r.font.size = Pt(6)


def _body_run(p, text: str, style: dict, *, bold=False, size=10, color=None):
    run = p.add_run(text)
    run.font.name = style["font"]
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color
    return run


def _contact_line(doc, cv: dict, style: dict):
    bits = [x for x in [cv.get("email"), cv.get("phone"), cv.get("location"),
                        cv.get("linkedin"), cv.get("website")] if x]
    if not bits:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _body_run(p, "  •  ".join(bits), style, size=10, color=RGBColor(0x55, 0x55, 0x55))


def _name_block(doc, cv: dict, style: dict):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _body_run(p, cv.get("full_name") or "Your Name", style,
              bold=True, size=style["name_size"], color=style["accent"])
    if cv.get("headline"):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _body_run(p2, cv["headline"], style, size=12, color=RGBColor(0x66, 0x66, 0x66))


def _summary(doc, cv: dict, style: dict):
    if not cv.get("summary"):
        return
    _heading(doc, "Summary", style)
    p = doc.add_paragraph()
    _body_run(p, cv["summary"], style, size=10)


def _experience(doc, cv: dict, style: dict):
    exps = cv.get("experience") or []
    if not exps:
        return
    _heading(doc, "Experience", style)
    for e in exps:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _body_run(p, (e.get("title") or "").strip(), style, bold=True, size=11)
        company_bits = [x for x in [e.get("company"), e.get("location")] if x]
        if company_bits:
            _body_run(p, "  —  " + " · ".join(company_bits), style, size=10,
                      color=RGBColor(0x55, 0x55, 0x55))
        dates = _date_range(e)
        if dates:
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(2)
            _body_run(p2, dates, style, size=9, color=RGBColor(0x88, 0x88, 0x88))
        for b in (e.get("bullets") or []):
            if not b:
                continue
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.space_after = Pt(0)
            _body_run(bp, b, style, size=10)


def _education(doc, cv: dict, style: dict):
    eds = cv.get("education") or []
    if not eds:
        return
    _heading(doc, "Education", style)
    for e in eds:
        p = doc.add_paragraph()
        _body_run(p, (e.get("degree") or "").strip(), style, bold=True, size=11)
        bits = [x for x in [e.get("school"), e.get("location")] if x]
        if bits:
            _body_run(p, "  —  " + " · ".join(bits), style, size=10,
                      color=RGBColor(0x55, 0x55, 0x55))
        dates = _date_range(e)
        if dates:
            p2 = doc.add_paragraph()
            _body_run(p2, dates, style, size=9, color=RGBColor(0x88, 0x88, 0x88))
        if e.get("details"):
            p3 = doc.add_paragraph()
            _body_run(p3, e["details"], style, size=10)


def _skills(doc, cv: dict, style: dict):
    s = cv.get("skills") or {}
    buckets = [
        ("Technical", s.get("technical") or []),
        ("Tools", s.get("tools") or []),
        ("Languages", s.get("languages") or []),
        ("Soft Skills", s.get("soft") or []),
    ]
    buckets = [(k, v) for k, v in buckets if v]
    if not buckets:
        return
    _heading(doc, "Skills", style)
    for label, items in buckets:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        _body_run(p, f"{label}: ", style, bold=True, size=10)
        _body_run(p, ", ".join(items), style, size=10)


def _certifications(doc, cv: dict, style: dict):
    certs = cv.get("certifications") or []
    if not certs:
        return
    _heading(doc, "Certifications", style)
    for c in certs:
        p = doc.add_paragraph(style="List Bullet")
        line = c.get("name") or ""
        if c.get("issuer"):
            line += f" — {c['issuer']}"
        if c.get("year"):
            line += f" ({c['year']})"
        _body_run(p, line, style, size=10)


def _projects(doc, cv: dict, style: dict):
    projs = cv.get("projects") or []
    if not projs:
        return
    _heading(doc, "Projects", style)
    for pr in projs:
        p = doc.add_paragraph()
        _body_run(p, pr.get("name") or "", style, bold=True, size=11)
        if pr.get("description"):
            p2 = doc.add_paragraph()
            _body_run(p2, pr["description"], style, size=10)


def _date_range(entry: dict) -> str:
    start, end = entry.get("start") or "", entry.get("end") or ""
    if start and end:
        return f"{start} – {end}"
    return start or end or ""


def _set_margins(doc: Document):
    for s in doc.sections:
        s.top_margin = Cm(1.6)
        s.bottom_margin = Cm(1.6)
        s.left_margin = Cm(1.8)
        s.right_margin = Cm(1.8)


def build_docx(cv: dict[str, Any], template: str = "minimal") -> bytes:
    style = STYLES.get(template, STYLES["minimal"])
    doc = Document()
    _set_margins(doc)

    _name_block(doc, cv, style)
    _contact_line(doc, cv, style)
    _summary(doc, cv, style)
    _experience(doc, cv, style)
    _education(doc, cv, style)
    _skills(doc, cv, style)
    _certifications(doc, cv, style)
    _projects(doc, cv, style)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
